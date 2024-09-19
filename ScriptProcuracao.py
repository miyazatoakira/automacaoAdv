import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import datetime
import os


local_dir = os.getcwd()



def proc_creator(nome_doc, nome, std_civil, ocupacao, cpf, rg, endereco, local, dia, mes, ano):
    doc = Document(fr'{local_dir}/modelo.docx')  # alterar caminho do arquivo, dependendo do dispositivo
    for style in doc.styles:
        if style.type == 1:
            style.font.size = Pt(12)
            style.font.name = 'Arial'

    title_paragraph = doc.add_paragraph("PROCURAÇÃO")
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_paragraph.runs[0].bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    justified_paragraph = doc.add_paragraph()
    run_contratante = justified_paragraph.add_run(f"Contratante: {nome}")
    run_contratante.bold = True
    justified_paragraph.add_run(f", {std_civil}, {ocupacao}, inscrito no cadastro de pessoas físicas CPF/MF sob o nº {cpf} e R.G. nº {rg}, residente e domiciliado na {endereco}.")
    justified_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    justified_paragraph = doc.add_paragraph()
    run_contratante = justified_paragraph.add_run("Contratada: IVONE PEREIRA DE SOUSA")
    run_contratante.bold = True
    justified_paragraph.add_run(", advogada, inscrita na OAB/SP, sob o número 437.365, Celular (11) 95072-8490 - E-mail: ivone.shiniti@adv.oabsp.org.br, com escritório na Rua João Mendes Júnior nº 41 – Centro – Francisco Morato – CEP 07910-220 – SP.")
    justified_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    justified_paragraph = doc.add_paragraph()
    run_contratante = justified_paragraph.add_run("PODERES:")
    run_contratante.bold = True
    justified_paragraph.add_run(' pelo presente instrumento o outorgante confere ao outorgado amplos poderes para o foro em geral, com cláusula "ad-judicia  et extra", em qualquer Juízo, Instância ou Tribunal, podendo propor contra quem de direito, as ações competentes e defendê-la nas contrárias, seguindo umas e outras, até final decisão, usando os recursos legais e acompanhando-os, conferindo-lhe ainda, poderes especiais para receber citação inicial, confessar, e conhecer a procedência do pedido, desistir, renunciar ao direito sobre que se funda a ação, transigir, firmar compromissos ou acordos, receber e dar quitação, podendo agir em Juízo ou fora dele, assim como  substabelecer esta a outrem, com ou sem reservas de iguais poderes, para agir em conjunto ou separadamente com o substabelecido. ')
    justified_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")
    doc.add_paragraph("")

    center_paragraph = doc.add_paragraph(f"{local}, {dia} de {mes} de {ano}")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    center_paragraph = doc.add_paragraph("______________________________________________")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    center_paragraph = doc.add_paragraph(nome)
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in center_paragraph.runs:
        run.bold = True

    doc.save(fr'{local_dir}/proc_{nome_doc}.docx')
    messagebox.showinfo("Sucesso", "Procuração criada com sucesso!")
