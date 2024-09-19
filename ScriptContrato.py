import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import datetime
import os


local_dir = os.getcwd()


def contrato_creator(nome_doc, nome, std_civil, ocupacao, cpf, rg, endereco, local, dia, mes, ano, valor_mensal, valor_total, valor_assinatura):
    doc = Document(fr'{local_dir}/modelo.docx')

    # Ajustando estilo da fonte
    for style in doc.styles:
        if style.type == 1:
            style.font.size = Pt(9)
            style.font.name = 'Calibri'

    title_paragraph = doc.add_paragraph("CONTRATO DE HONORÁRIOS ADVOCATÍCIOS")
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_paragraph.runs[0].bold = True

    doc.add_paragraph("")
    justified_paragraph = doc.add_paragraph()
    run_contratante = justified_paragraph.add_run(f"Pelo presente instrumento de contrato de honorários advocatícios, a advogada IVONE PEREIRA DE SOUSA, OAB/SP 437.365, com escritório na Rua João Mendes Júnior, nº 41 1º Andar sala 01- Centro – Francisco Morato – São Paulo/SP, ora contratante – {nome}, {std_civil}, {ocupacao}, inscrito no cadastro de pessoas físicas CPF/MF sob o nº {cpf} e R.G. nº {rg}, residente e domiciliado na {endereco}, convencionam e contratam o seguinte:")
    justified_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")
    doc.add_paragraph(f"Cláusula 1a. A Advogada contratada obriga-se, face ao mandato judicial a que lhe foi outorgado, a prestar seus serviços profissionais, na defesa dos direitos do contratante, nas instâncias competentes, administrativas e judiciais.")

    doc.add_paragraph(f"Cláusula 2a. A remuneração pelos serviços prestados será no valor de R$ {valor_total}, os quais serão pagos da seguinte forma: R$ {valor_assinatura} na assinatura deste, e o restante será pago em 5 parcelas de R$ {valor_mensal}, com vencimento no dia 20 de cada mês subsequente.")

    doc.add_paragraph(f"Cláusula 3a. O Contratante obriga-se a fornecer todos os documentos, informações e meios de prova necessários ao bom desempenho profissional dos contratados na defesa da lide, bem como adiantar valores referentes à satisfação de custas judiciais e extrajudiciais. Obriga-se inclusive, a pagar as despesas decorrentes de transportes e viagens, quando estes se fizerem necessários fora do domicílio do contratado, bem como honorários com advogado correspondente, quando for o caso;")

    doc.add_paragraph(f"Cláusula 4a. Serão pagas pelo contratante as despesas com perícias que se fizerem necessárias durante a tramitação do processo, inclusive as contábeis, por conta da realização de cálculos determinados pelo Juízo;")

    doc.add_paragraph(f"Cláusula 5a. Convencionam que os honorários advocatícios poderão ser exigidos imediatamente se houver conciliação entre as partes em litígio, bem como no caso de não haver prosseguimento da ação por interesse do contratante ou, se for revogado o mandato sem que os contratados tenham concorrido com culpa ou dolo. Para efeito de aplicação do percentual convencionado na cláusula 2a, será considerado o valor total da liquidação de todos os pedidos da inicial ou o valor total da liquidação;")

    doc.add_paragraph(f"Cláusula 6a. Os honorários devidos corrigir-se-ão na data do respectivo pagamento, de acordo com os índices oficiais, podendo ser exigidos integralmente no recebimento parcial do processo;")

    doc.add_paragraph(f"Cláusula 7a. Caso ocorra desistência da ação por parte do reclamante, será devido aos advogados, a título de pagamento dos serviços realizados, o montante dos serviços prestados até o momento. Será considerada desistência da ação o não comparecimento do autor na audiência em que era imprescindível a sua presença;")

    doc.add_paragraph(f"Clausula 8ª. Fica convencionado que o foro para dirimir dúvidas decorrentes do presente contrato de honorários será o da Comarca de Francisco Morato, inclusive para execução dos honorários advocatícios.")

    doc.add_paragraph(f"Clausula 9ª. Serviços de deslocamento do Advogado, tais como Gasolina e Alimentação serão cobrados ao final do Processo.")

    doc.add_paragraph("")
    center_paragraph = doc.add_paragraph(f"{local}, {dia} de {mes} de {ano}")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    center_paragraph = doc.add_paragraph("______________________________________________")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    center_paragraph = doc.add_paragraph(nome)
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    center_paragraph = doc.add_paragraph("______________________________________________")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    center_paragraph = doc.add_paragraph("IVONE PEREIRA DE SOUSA")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(fr'{local_dir}/contrato_{nome_doc}.docx') # alterar caminho do arquivo, dependendo do dispositivo
    messagebox.showinfo("Sucesso", "Contrato de honorários advocatícios criado com sucesso!")
