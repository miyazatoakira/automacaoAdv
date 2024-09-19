
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import datetime
import os


local_dir = os.getcwd()


def hipo_creator(nome_doc, nome, std_civil, ocupacao, cpf, rg, endereco, local, dia, mes, ano):
    doc = Document(fr'{local_dir}/modelo.docx')
    for style in doc.styles:
        if style.type == 1:
            style.font.size = Pt(12)
            style.font.name = 'Arial'

    title_paragraph = doc.add_paragraph("DECLARAÇÃO DE HIPOSSUFICIÊNCIA")
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_paragraph.runs[0].bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    justified_paragraph = doc.add_paragraph()
    run_contratante = justified_paragraph.add_run("Contratante: ")
    run_contratante.bold = True
    run_nome = justified_paragraph.add_run(nome)
    run_nome.bold = True
    justified_paragraph.add_run(f", {std_civil}, {ocupacao}, inscrito no cadastro de pessoas físicas CPF/MF sob o nº {cpf} e R.G. nº {rg}, residente e domiciliado na {endereco}, declara não possuir condições financeiras para pagamento das custas do processo, honorários periciais e demais despesas processuais, sem prejuízo de seu sustento e de sua família, por ser pobre na acepção jurídica do termo em consonância com o art. 791-A da CLT.")
    justified_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    justified_paragraph = doc.add_paragraph('                                                                  Responsabiliza-se pela veracidade das informações constantes nesta declaração, sob pena das cominações legais respectivas, sujeitando-se às sanções administrativas, civis e criminais, nos termos dos artigos 2° e 3°, ambos da lei n° 7.115/83.')
    justified_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")
    doc.add_paragraph("")

    center_paragraph = doc.add_paragraph(f"{local}, {dia} de {mes} de {ano}")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    center_paragraph = doc.add_paragraph("______________________________________________")
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    center_paragraph = doc.add_paragraph(nome)
    center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in center_paragraph.runs:
        run.bold = True

    doc.save(fr'{local_dir}/hipo_{nome_doc}.docx')
    messagebox.showinfo("Sucesso", "Declaração de Hipossuficiência criada com sucesso!")
